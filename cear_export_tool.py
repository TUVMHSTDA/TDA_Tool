import streamlit as st
import docx
import time
from pdf_extractor import extract_txt_list  # Daniel's script
from export_data import main as export_td
import fitz  # pymupdf

class CEARExportTool:
    def __init__(self):
        if "word_file_path" not in st.session_state:
            st.session_state.word_file_path = ""
        if "cear_template_path" not in st.session_state:
            st.session_state.cear_template_path = ""
        if "output_path" not in st.session_state:
            st.session_state.output_path = ""

    def run(self):
        st.title("CEAR Export Tool")
        st.write("Upload the completed TDAR document and CEAR template to proceed with the export.")

        uploaded_word_file = st.file_uploader("Upload Completed Word Document", type=["docx"])
        if uploaded_word_file is not None:
            st.session_state.word_file_path = f"/tmp/{uploaded_word_file.name}"
            with open(st.session_state.word_file_path, "wb") as f:
                f.write(uploaded_word_file.getbuffer())
        st.text_input("Completed Word Document Directory", value=st.session_state.word_file_path, disabled=True)

        uploaded_cear_template = st.file_uploader("Upload CEAR Template", type=["docx"])
        if uploaded_cear_template is not None:
            st.session_state.cear_template_path = f"/tmp/{uploaded_cear_template.name}"
            with open(st.session_state.cear_template_path, "wb") as f:
                f.write(uploaded_cear_template.getbuffer())
        st.text_input("CEAR Template Directory", value=st.session_state.cear_template_path, disabled=True)

        if st.button("Export"):
            self.export_with_progress()

    def export_with_progress(self):
        progress_bar = st.progress(0)
        for i in range(1, 10):
            progress_bar.progress(i * 10)
            st.write(f"Step {i}/9: Processing...")
            time.sleep(0.5)
        
        # Start document processing
        self.process_documents()
        st.success("Export complete!")
        
        # Auto-download button for the output document
        with open(st.session_state.output_path, "rb") as file:
            st.download_button(
                label="Download Completed Document",
                data=file,
                file_name="Evaluation_Assessment_Report_Filled.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    def process_documents(self):
        # Load the documents
        source_doc = docx.Document(st.session_state.word_file_path)
        template_doc = docx.Document(st.session_state.cear_template_path)

        # Remove hidden text from the source document
        remove_hidden_text(source_doc.element)

        # Remove and replace all content controls in the source document
        for p in source_doc.paragraphs:
            replace_content_control(p._element)
        
        for t in source_doc.tables:
            replace_content_control(t._element)

        # Extract all text from the source document
        full_text = "\n".join([para.text for para in source_doc.paragraphs])

        # Define placeholders and markers within the template document
        sections = [
            ('%productname%', "Product or trade name:", "Reference list of applicable or corresponding UDI-DIs assigned: "),
            ('%basicudi%', "Basic UDI device identifier(s)", "Test specifications"),
            ('%mdan%', "MDN / MDA and MDS code(s)", "Basic UDI device identifier(s)"),
            ('%manufacturer%', "Legal manufacturer", "Single registration number"),
            ('%srn%', "Single registration number", "Authorised representative"),
            ('%authorep%', "Authorised representative", "Production location(s)"),
            ('%devicedescription%', "General description of the device:", 
             "Intended purpose of the device in accordance with the current instructions for use:"),
            ('%intendedpurpose%', "Intended purpose of the device in accordance with the current instructions for use: ", 
             "Intended user of the device:"),
            ('%intendeduser%', "Intended user of the device:", 
             "Intended patient population, patient selection criteria, indications, contraindications, warnings (MDR Annex II Section 1.1(c))"),
            ('%populationcontra%', "Intended patient population, patient selection criteria, indications, contraindications, warnings (MDR Annex II Section 1.1(c))", 
             "Warning:"),
            ('%contraindications%', "Contraindication:", "Warning:"),
            ('%warnings%', "Warning:", "Principles of operation of the device and its mode of action (MDR Annex II, Section 1.1(d))"),
            ('%deviceconfig%', "List of all configurations and variants of the device intended to be made available on the market:", 
             "Accessories and device combinations (MDR Annex II Sections 1.1(h), 1.1(i))"),
            ('%accessories%', "List of accessories provided with the device:", "Combined devices"),
            ('%previousgen%', "Overview of a previous generation or generations of the device: ", "INFORMATION SUPPLIED BY THE MANUFACTURER ")
        ]

        # Replace the placeholders in the template document
        for placeholder, start_marker, end_marker in sections:
            fill_placeholder(template_doc, placeholder, full_text, start_marker, end_marker)

        # Save the filled template
        output_path = "/tmp/Evaluation_Assessment_Report_Filled.docx"
        template_doc.save(output_path)
        st.session_state.output_path = output_path
        st.write("Document saved successfully!")

# Utility functions integrated from auto_populate.py
def replace_content_control(element, namespaces=None):
    kwargs = {} if not namespaces else {"namespaces": namespaces}
    sdt_elements = element.xpath(".//w:sdt", **kwargs)
    for e in sdt_elements:
        inner_elements = e.xpath(".//w:sdt", namespaces=e.nsmap)
        for inner_element in inner_elements:
            replace_content_control(inner_element, namespaces=e.nsmap)
        content_elements = e.xpath(".//w:sdtContent/*", namespaces=e.nsmap)
        for c in content_elements:
            e.addprevious(c)
        parent = e.getparent()
        if parent is not None:
            parent.remove(e)

def remove_hidden_text(element, namespaces=None):
    kwargs = {} if not namespaces else {"namespaces": namespaces}
    hidden_elements = element.xpath(".//w:p[w:pPr/w:pStyle[@w:val='Informationinvisibelblau']]", **kwargs) + \
                      element.xpath(".//w:p[w:pPr/w:pStyle[@w:val='Informationinvisibelgrn']]", **kwargs)
    for e in hidden_elements:
        parent = e.getparent()
        if parent is not None:
            parent.remove(e)

def extract_text_between_markers(text, start_marker, end_marker):
    try:
        start_index = text.index(start_marker) + len(start_marker)
        end_index = text.index(end_marker)
        return "\n" + text[start_index:end_index].strip() + "\n"
    except ValueError:
        return "Markers not found in the provided text."

def fill_placeholder(template_doc, placeholder, full_text, start_marker, end_marker):
    extracted_text = extract_text_between_markers(full_text, start_marker, end_marker)
    for para in template_doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, extracted_text)
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, extracted_text)

if __name__ == "__main__":
    tool = CEARExportTool()
    tool.run()
