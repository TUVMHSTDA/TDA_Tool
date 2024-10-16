import json
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

def load_json(json_path):
    # Load JSON data from a given path
    try:
        with open(json_path, 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"Error: The file {json_path} was not found.")
        return None
    except json.JSONDecodeError:
        print(f"Error: The file {json_path} is not a valid JSON file.")
        return None

def check_multiple_selections(data):
    
    previous = ""
    current = ""
    multiple = False
    existing = []

    try:

        for entry in data:
            if entry["Selection"] == True:
                # Get all the selected fields
                existing.append(entry["Sub-Section"])

        # Sort to ensure that fields of same name is grouped together; easier to compare
        existing.sort()
        print(f"Selected fields: {existing}")

        for field in existing:
            current = field
            
            # check that there are no duplicates by comparing against previous field
            if current != previous:
                print("yay")
                previous = current
                multiple = False
            else:
                print("You have selected multiple options of the same field. Please ensure that only one per field is selected.")
                multiple = True
                break
    except:
        print("no")

    return multiple

def create_placeholders(data):
    # Create a dictionary to map placeholders to their corresponding answers
    placeholders = {
        '{%PROJECT_SOLD_TO%}': '',
        '{%PROJECT_SOLD_TO_ADR%}': '',
        '{%SRN_Manufacturer%}': '',
        '{%PROJECT_CBW_EU_NAME%}': '',
        '{%PROJECT_CBW_PRODUCT%}': '',
        '{%EMDN%}': '',
        '{%Basic UDI-DI(s)%}': ''
    }

    if not data:
        return placeholders

    # Populate the placeholders dictionary
    try:
        for entry in data:
            if entry["Selection"] == True:
                if entry['Sub-Section'] == 'Legal Manufacturer':
                    placeholders['{%PROJECT_SOLD_TO%}'] = entry["Extracted Data Points"]
                elif entry['Sub-Section'] == 'Product Location(s)':
                    placeholders['{%PROJECT_SOLD_TO_ADR%}'] = entry["Extracted Data Points"]
                elif entry['Sub-Section'] == 'Single Registration Number':
                    placeholders['{%SRN_Manufacturer%}'] = entry["Extracted Data Points"]
                elif entry['Sub-Section'] == 'Authorized Representative':
                    placeholders['{%PROJECT_CBW_EU_NAME%}'] = entry["Extracted Data Points"]
                elif entry['Sub-Section'] == 'Test Subject':
                    placeholders['{%PROJECT_CBW_PRODUCT%}'] = entry["Extracted Data Points"]
                elif entry['Sub-Section'] == 'EMDN Code':
                    placeholders['{%EMDN%}'] = entry["Extracted Data Points"]
                elif entry['Sub-Section'] == 'Basic UDI-Device Identifier':
                    placeholders['{%Basic UDI-DI(s)%}'] = entry["Extracted Data Points"]


    except KeyError as e:                     
        print(f"Error: Missing expected key in JSON data - {e}")

    return placeholders

def replace_placeholders(doc_path, placeholders, output_path):
    # Replace placeholders in the document with corresponding values
    try:
        doc = Document(doc_path)
    except FileNotFoundError:
        print(f"Error: The file {doc_path} was not found.")
        return
    except PackageNotFoundError:
        print(f"Error: The file {doc_path} is not a valid Word document.")
        return

    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    try:
        doc.save(output_path)
        print("Export complete!")
    except Exception as e:
        print(f"Error: Could not save the document to {output_path} - {e}")

def main():
    # json_path = input("Json file path: ").strip().replace('"', '').replace("'", "")
    json_path = r"./edited_data.json"
    # r"C:\Users\iande-da\Desktop\Notes\technical_docs\output\DOC\output.json"

    # doc_path = input("Word doc template path: ").strip().replace('"', '').replace("'", "")
    doc_path = r"./11192_MDR Technical Documentation Assessment Report_Rev14_MDA0315.docx"
    # r"C:\Users\iande-da\Desktop\Notes\technical_docs\11192_MDR Technical Documentation Assessment Report_Rev14_MDA0315.docx"
    
    # output_path = input("Output path: ").strip().replace('"', '').replace("'", "")
    output_path = r"./exported.docx"
    # r'C:\Users\iande-da\Desktop\Notes\technical_docs\exported.docx'
    
    data = load_json(json_path)     # returns json as dict
    selection = check_multiple_selections(data)     # checks if there are duplicates in the table
    if selection == False:
        placeholders = create_placeholders(data)
        replace_placeholders(doc_path, placeholders, output_path)


if __name__ == "__main__":
    main()
